from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import  Pt #磅数
from docx.oxml.ns import qn #中文格式
import  time

price=input('请输入今日价格')
company_list=['佟丽娅','汪小敏','赵丽颖','杨紫','唐嫣','张天爱','高圆圆']
today1=time.strftime('%Y{y}%m{m}%d{d}',time.localtime()).format(y='年',m='月',d='日') #中文方式输出
print(today1)

for i in company_list:
    document=Document()
    # 设置文档的基础字体
    document.styles['Normal'].font.name=u'宋体'
    #支持中文设置的字体
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
    #初始化第一个自然段落
    p1=document.add_paragraph()
    #对齐方式为剧中，没有这句话默认左对齐
    p1.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    run1=p1.add_run('关于下达%s产品价格的通知' % (today1))
    #西文字体
    run1.font.name='微软雅黑'
    #中文字体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
    # 设置为21磅
    run1.font.size=Pt(21)
    # 设置加粗
    run1.font.bold=True
    p1.space_after=Pt(5)
    p1.space_befor = Pt(5)
    #_________________________________
    p2 = document.add_paragraph()
    # 对齐方式为剧中，没有这句话默认左对齐
    # p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #默认左对齐，不用设置
    run2 = p2.add_run(i +':')
    # 西文字体
    run2.font.name = '仿宋_GB2312'
    # 中文字体
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    # 设置为21磅
    run2.font.size = Pt(16)
    # 设置加粗
    run2.font.bold = True

    #_________________________________
    p3 = document.add_paragraph()
    run3 = p3.add_run("      根据公司安排，为提供优质客户服务，我单位拟定了今天黄金价格为%s元，特此通知。" % price)
    # 西文字体
    run3.font.name = '仿宋_GB2312'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run3.font.size = Pt(16)
    run3.font.bold = True
    #_________________________________
    p4 = document.add_paragraph()
    run4 = p4.add_run('联系人：刘德华    电话：188888888888')
    # 西文字体
    run4.font.name = '仿宋_GB2312'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    run4.font.size = Pt(16)
    run4.font.bold = True


    document.save('%s-价格通知.docx' % i )