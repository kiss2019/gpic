from docx import Document
from docx.oxml.ns import qn  # 中文格式
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt  # 磅数
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 对齐
from docx.shared import Inches #图片尺寸
import time

price = input('输入今天的价格：')
company_list = ['刘德华', '张学友', '黎明', '倪妮', '汪小敏', '同仁堂']
today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日') #time对中文支付不好，要这样写
print(today)

for i in company_list:
    document = Document()
    # 设置文档的基础字体
    document.styles['Normal'].font.name = u'宋体'
    # 只有加了上面这行中文字体才有效果（库对中文支持不好）
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')

    document.add_picture('123.jpg',width=Inches(6))
    # document.add_page_break()
    p1 = document.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置对齐方式为居中，没有这句话默认左对齐
    run1 = p1.add_run('关于下达%s产品价格的通知' % (today))
    #设置西文字体
    run1.font.name = '微软雅黑'
    #设置中文体
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    run1.font.size = Pt(21) #字体磅数
    run1.font.bold = True
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)
#---------------------------------------------
    p2 = document.add_paragraph()
      # 设置对齐默认左对齐
    run2 = p2.add_run(i + '：')
    # 设置西文字体
    run2.font.name = '仿宋_GB2312'
    # 设置中午体
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB231')
    run2.font.size = Pt(16)  # 字体磅数
    run2.font.bold = True

 # ---------------------------------------------
    p3 = document.add_paragraph()
    # 设置对齐默认左对齐
    run3 = p3.add_run('    根据公司安排，为提供优质客户服务，我单位拟定了今日黄金价格{}元，特此通知。'.format(price))
    # 设置西文字体
    run3.font.name = '仿宋_GB2312'
    # 设置中午体
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB231')
    run3.font.size = Pt(16)  # 字体磅数
    run3.font.bold = True
 # ---------------------------------------------


    table = document.add_table(rows=3,cols=3,style='Table Grid')
    table.cell(0, 0).merge(table.cell(0, 2)) #合并0行0列到0行2列
    table_run1=table.cell(0,0).paragraphs[0].add_run("xx产品的报价表")
    table_run1.font.name=u'隶书'
    table_run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB231')
    table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #居中
    table.cell(1, 0).text='日期'
    table.cell(1, 1).text = '价格'
    table.cell(1, 2).text = '备注'
    table.cell(2, 0).text = today
    table.cell(2, 1).text = str(price)
    table.cell(2, 2).text = ''
#------------------------------------
    p4 = document.add_paragraph()
    # 设置对齐默认左对齐
    run4 = p4.add_run('           (联系人：小优     电话：13999999999)')
    # 设置西文字体
    run4.font.name = '仿宋_GB2312'
    # 设置中午体
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB231')
    run4.font.size = Pt(16)  # 字体磅数
    run4.font.bold = True

    document.add_page_break()

    p5 = document.add_paragraph()
    run5 = p4.add_run('这狗日的是广告')

    document.save('%s-价格通知2.docx' % i)
print(document.paragraphs[1].style.font.name)
print(document.paragraphs[1].style.font.size)
